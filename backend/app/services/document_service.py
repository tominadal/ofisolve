import os
import datetime
from typing import List, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from loguru import logger

from app.models.db_models import DocumentoLibreria, Cliente, Tramite
from app.core.config import get_settings

class DocumentService:
    """
    Servicio profesional para la gestión de documentos en el filesystem local y base de datos.
    Sincroniza la persistencia física con el registro de metadatos.
    """

    def __init__(self):
        settings = get_settings()
        self.base_uploads = os.path.join(os.getcwd(), "uploads")
        self.clientes_dir = os.path.join(self.base_uploads, "clientes")
        os.makedirs(self.clientes_dir, exist_ok=True)

    async def listar_por_cliente(self, db: AsyncSession, cliente_id: int) -> List[dict]:
        """Lista todos los documentos registrados para un cliente."""
        stmt = select(DocumentoLibreria).where(DocumentoLibreria.cliente_id == cliente_id)
        result = await db.execute(stmt)
        docs = result.scalars().all()
        return [
            {
                "id": d.id,
                "nombre": d.nombre,
                "tipo": d.tipo,
                "path": d.path,
                "fecha_subida": d.fecha_subida,
                "tramite_id": d.tramite_id
            }
            for d in docs
        ]

    async def obtener_contenido(self, db: AsyncSession, doc_id: int) -> Optional[str]:
        """Lee el contenido de un archivo de texto/docx registrado."""
        stmt = select(DocumentoLibreria).where(DocumentoLibreria.id == doc_id)
        result = await db.execute(stmt)
        doc = result.scalars().first()
        
        if not doc:
            return None
            
        import sys
        from pathlib import Path
        # Determinar el BASE_DIR asumiendo que el archivo está en backend/app/services/document_service.py
        BASE_DIR = Path(__file__).parent.parent.parent.parent
        
        # Si la ruta guardada es absoluta, se usa. Si es relativa, se anexa a BASE_DIR
        doc_path = doc.path
        if not os.path.isabs(doc_path):
            doc_path = os.path.join(str(BASE_DIR), doc_path)
            
        if not os.path.exists(doc_path):
            logger.warning(f"Archivo no encontrado en disco: {doc_path}")
            return f"[Error: El archivo físico '{doc.nombre}' no se encuentra en el servidor. Contacte al administrador.]"

        # Por ahora solo soportamos lectura de archivos de texto/planos para el LLM
        # En una versión pro extenderíamos con PDFMiner o similar
        try:
            with open(doc_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error leyendo archivo {doc.path}: {e}")
            return f"[Error: No se pudo leer el contenido del archivo {doc.nombre}]"

    async def guardar_contenido(self, db: AsyncSession, doc_id: int, contenido: str) -> bool:
        """Actualiza el contenido de un archivo registrado."""
        stmt = select(DocumentoLibreria).where(DocumentoLibreria.id == doc_id)
        result = await db.execute(stmt)
        doc = result.scalars().first()
        
        if not doc:
            logger.error(f"Intento de guardar documento inexistente: ID {doc_id}")
            return False
            
        import sys
        from pathlib import Path
        BASE_DIR = Path(__file__).parent.parent.parent.parent
        
        doc_path = doc.path
        if not os.path.isabs(doc_path):
            doc_path = os.path.join(str(BASE_DIR), doc_path)
            
        if not os.path.exists(doc_path):
            logger.error(f"El archivo físico {doc_path} no existe, se creará uno nuevo.")
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)
            
        try:
            with open(doc_path, "w", encoding="utf-8") as f:
                f.write(contenido)
            logger.info(f"Documento sobrescrito correctamente: {doc_path}")
            # Si queremos actualizar una fecha de actualización en DB:
            # doc.fecha_actualizacion = datetime.datetime.utcnow()
            # await db.commit()
            return True
        except Exception as e:
            logger.error(f"Error guardando archivo {doc_path}: {e}")
            return False

    async def guardar_documento(
        self, 
        db: AsyncSession, 
        nombre: str, 
        contenido: str, 
        workspace_id: int,
        cliente_id: Optional[int] = None,
        tramite_id: Optional[int] = None
    ) -> DocumentoLibreria:
        """Guarda un nuevo documento en el FS y lo registra en la DB."""
        
        # Determinar ruta
        target_dir = self.clientes_dir
        if cliente_id:
            target_dir = os.path.join(target_dir, f"cliente_{cliente_id}")
        if tramite_id:
            target_dir = os.path.join(target_dir, f"tramite_{tramite_id}")
        
        os.makedirs(target_dir, exist_ok=True)
        
        # Evitar colisión de nombres
        base_name, ext = os.path.splitext(nombre)
        if not ext: ext = ".txt"
        final_path = os.path.join(target_dir, f"{base_name}{ext}")
        
        counter = 1
        while os.path.exists(final_path):
            final_path = os.path.join(target_dir, f"{base_name}_{counter}{ext}")
            counter += 1

        # Escribir en disco
        with open(final_path, "w", encoding="utf-8") as f:
            f.write(contenido)
        
        logger.info(f"Documento guardado en disco: {final_path}")

        # Registrar en DB
        nuevo_doc = DocumentoLibreria(
            workspace_id=workspace_id,
            cliente_id=cliente_id,
            tramite_id=tramite_id,
            nombre=os.path.basename(final_path),
            tipo=ext.replace(".", ""),
            path=final_path,
            fecha_subida=datetime.datetime.utcnow()
        )
        db.add(nuevo_doc)
        await db.commit()
        await db.refresh(nuevo_doc)
        
        return nuevo_doc

    def generar_docx(self, datos_finales: dict, nombre_plantilla: str) -> "pathlib.Path":
        """
        Genera un archivo .docx basado en un diccionario de datos.
        Retorna la ruta al archivo generado.
        """
        import pathlib
        from docx import Document
        from docx.shared import Pt
        from bs4 import BeautifulSoup
        
        doc = Document()
        
        # Titulo
        titulo = doc.add_heading(datos_finales.get("tipo_certificacion_display", "DOCUMENTO NOTARIAL"), 0)
        titulo.alignment = 1 # Center
        
        # Escribano info
        p = doc.add_paragraph()
        p.add_run(f"Escribano/a: {datos_finales.get('nombre_escribano', 'N/A')}").bold = True
        p.add_run(f" | Registro N°: {datos_finales.get('nro_registro', 'N/A')}")
        
        # Contenido (parseamos un poco el HTML de Quill si es necesario, pero como fallback guardamos texto crudo)
        texto_raw = datos_finales.get("texto_certificacion", "")
        
        try:
            soup = BeautifulSoup(texto_raw, "html.parser")
            texto_limpio = soup.get_text(separator="\n")
        except:
            texto_limpio = texto_raw
            
        doc.add_paragraph(texto_limpio)
        
        # Guardar archivo
        import time
        safe_name = f"documento_generado_{int(time.time())}.docx"
        target_dir = os.path.join(self.base_uploads, "generados")
        os.makedirs(target_dir, exist_ok=True)
        final_path = os.path.join(target_dir, safe_name)
        
        doc.save(final_path)
        logger.info(f"Documento DOCX generado en: {final_path}")
        return pathlib.Path(final_path)
