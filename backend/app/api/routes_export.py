import os
import io
from fastapi import APIRouter, HTTPException, Response
from pydantic import BaseModel
from docx import Document
from fpdf import FPDF
from loguru import logger

router = APIRouter(tags=["Exportación"])

class ExportRequest(BaseModel):
    titulo: str
    contenido: str
    formato: str # 'docx' o 'pdf'

@router.post("")
@router.post("/")
async def exportar_documento(payload: ExportRequest):
    """
    Genera y devuelve un archivo en el formato solicitado.
    """
    try:
        if payload.formato == "docx":
            return await generate_docx(payload.titulo, payload.contenido)
        elif payload.formato == "pdf":
            return await generate_pdf(payload.titulo, payload.contenido)
        else:
            raise HTTPException(status_code=400, detail="Formato no soportado")
    except Exception as e:
        logger.error(f"Error exportando: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error al generar archivo: {str(e)}")

async def generate_docx(titulo: str, contenido: str):
    doc = Document()
    doc.add_heading(titulo, 0)
    
    # El contenido puede venir con saltos de línea \n
    paragraphs = contenido.split('\n')
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p)
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return Response(
        content=file_stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={titulo.replace(' ', '_')}.docx"
        }
    )

async def generate_pdf(titulo: str, contenido: str):
    pdf = FPDF()
    pdf.add_page()
    
    # Añadir fuente que soporte caracteres latinos if needed, pero fpdf2 tiene una básica
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, titulo, ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Helvetica", size=12)
    # Reemplazar caracteres que puedan romper fpdf si no se usa unicode
    safe_content = contenido.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.multi_cell(0, 10, safe_content)
    
    file_stream = io.BytesIO()
    # fpdf2 output() can return bytes
    pdf_bytes = pdf.output()
    
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={titulo.replace(' ', '_')}.pdf"
        }
    )
